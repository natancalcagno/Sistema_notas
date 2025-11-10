from django.shortcuts import render, redirect, get_object_or_404
from django.views.generic import View, ListView, CreateView, UpdateView, DeleteView, TemplateView
from django.contrib.auth.views import LoginView
from django.urls import reverse_lazy
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.contrib import messages
from django.http import HttpResponse, FileResponse, JsonResponse
from django.template.loader import render_to_string
from django.db.models import Sum, Count, Q
from django.utils import timezone
from django.core.paginator import Paginator
from django.views.decorators.cache import cache_page
from django.utils.decorators import method_decorator
from django.core.exceptions import ValidationError
from datetime import datetime, timedelta
from django.conf import settings
# Importação resiliente de python-docx para evitar crash em ambientes sem o pacote
try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except Exception:
    Document = None
    Pt = Cm = Inches = RGBColor = None
    WD_ALIGN_PARAGRAPH = None
    def qn(x):
        return x
    DOCX_AVAILABLE = False
import os
import io
from .models import Contrato, Nota, Usuario, TokenRedefinicaoSenha
from .forms import ContratoForm, NotaForm, UsuarioForm, UsuarioUpdateForm, AlterarSenhaForm, EsqueciSenhaForm, RedefinirSenhaForm
from .views_password import AlterarSenhaView, EsqueciSenhaView, RedefinirSenhaView
from .mixins import AdminRequiredMixin, OwnerRequiredMixin, MessageMixin, FilterMixin, PaginationMixin, SearchMixin
from .services import ContratoService, NotaService, UsuarioService, DashboardService, RelatorioService
from .cache_utils import CacheManager
from .pagination import OptimizedPaginator

# Views de Autenticação
class LoginView(LoginView):
    template_name = 'login.html'
    fields = '__all__'
    redirect_authenticated_user = True
    
    def get_success_url(self):
        return reverse_lazy('core:home')

    def post(self, request, *args, **kwargs):
        """
        Em ambiente sem banco (ENGINE dummy), evita 500 ao tentar autenticar.
        Mostra uma mensagem amigável e mantém na página de login.
        """
        try:
            from django.db import connections
            engine = connections['default'].settings_dict.get('ENGINE', '')
        except Exception:
            engine = ''

        if engine == 'django.db.backends.dummy':
            messages.error(request, 'Autenticação indisponível: banco de dados não configurado. Defina DATABASE_URL ou variáveis DB_* no ambiente de produção.')
            return self.get(request, *args, **kwargs)

        return super().post(request, *args, **kwargs)

# View da Home
class HomeView(LoginRequiredMixin, TemplateView):
    template_name = 'dashboard.html'
    
    def get(self, request, *args, **kwargs):
        # Verificar se é uma requisição AJAX
        if request.GET.get('ajax'):
            return self.get_ajax_response(request)
        # Evitar consultas ao banco em ambiente dummy
        try:
            from django.db import connections
            engine = connections['default'].settings_dict.get('ENGINE', '')
        except Exception:
            engine = ''
        if engine == 'django.db.backends.dummy':
            messages.warning(request, 'Dashboard em modo leitura: banco de dados não configurado. Algumas métricas não serão exibidas.')
        return super().get(request, *args, **kwargs)
    
    def get_ajax_response(self, request):
        """Responder a requisições AJAX para carregar notas"""
        try:
            nota_service = NotaService(request.user)
            
            # Obter filtros da requisição
            filtros = {
                'empresa': request.GET.get('empresa'),
                'setor': request.GET.get('setor'),
                'data_inicio': request.GET.get('data_inicio'),
                'data_fim': request.GET.get('data_fim'),
                'empenho': request.GET.get('empenho')
            }
            
            # Remover filtros vazios
            filtros = {k: v for k, v in filtros.items() if v}
            
            # Obter notas
            notas = nota_service.listar_notas(filtros)
            
            # Paginação
            page = int(request.GET.get('page', 1))
            paginator = Paginator(notas, 15)
            page_obj = paginator.get_page(page)
            
            # Renderizar HTML das notas
            from django.template.loader import render_to_string
            html = render_to_string('partials/notas_list.html', {
                'notas': page_obj.object_list
            })
            
            # Renderizar paginação
            pagination_html = render_to_string('partials/pagination.html', {
                'page_obj': page_obj
            })
            
            return JsonResponse({
                'success': True,
                'html': html,
                'pagination': pagination_html,
                'count': len(page_obj.object_list),
                'total': paginator.count
            })
            
        except Exception as e:
            return JsonResponse({
                'success': False,
                'error': str(e)
            })
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Evitar consultas ao banco em ambiente dummy
        try:
            from django.db import connections
            engine = connections['default'].settings_dict.get('ENGINE', '')
        except Exception:
            engine = ''
        if engine == 'django.db.backends.dummy':
            context['estatisticas'] = {}
            context['graficos'] = {}
            context['notas'] = []
            context['empresas'] = []
            return context

        # Usar service para obter estatísticas
        dashboard_service = DashboardService(self.request.user)
        
        try:
            # Obter estatísticas gerais
            context['estatisticas'] = dashboard_service.obter_estatisticas_gerais()
            
            # Obter dados para gráficos
            context['graficos'] = dashboard_service.obter_dados_graficos()
            
            # Adicionar lista de empresas do cache
            context['empresas'] = CacheManager.get_empresas_list()
            
            # Carregar notas iniciais para o dashboard
            nota_service = NotaService(self.request.user)
            notas = nota_service.listar_notas()
            
            # Paginação inicial
            paginator = Paginator(notas, 15)
            page_obj = paginator.get_page(1)
            context['notas'] = page_obj.object_list
            context['page_obj'] = page_obj
            
        except Exception as e:
            messages.error(self.request, f'Erro ao carregar dashboard: {str(e)}')
            context['estatisticas'] = {}
            context['graficos'] = {}
            context['notas'] = []
        
        return context

# Views de Contratos
class ContratoListView(LoginRequiredMixin, FilterMixin, PaginationMixin, ListView):
    model = Contrato
    template_name = 'lista_contratos.html'
    context_object_name = 'contratos'
    paginate_by = 10

    def get_queryset(self):
        contrato_service = ContratoService(self.request.user)
        
        # Obter filtros da requisição
        filtros = {
            'empresa': self.request.GET.get('empresa'),
            'numero': self.request.GET.get('numero'),
            'status': self.request.GET.get('status'),
            'vencendo': self.request.GET.get('vencendo') == 'true',
            'data_inicio': self.request.GET.get('data_inicio'),
            'data_termino': self.request.GET.get('data_termino')
        }
        
        # Remover filtros vazios
        filtros = {k: v for k, v in filtros.items() if v}
        
        return contrato_service.listar_contratos(filtros)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Obter contratos que estão vencendo (30 dias ou menos)
        from datetime import date
        contratos_alerta = Contrato.objects.filter(
            data_termino__gte=date.today()
        ).extra(
            where=["(data_termino - %s) <= alerta_vencimento"],
            params=[date.today()]
        ).order_by('data_termino')
        
        context['contratos_alerta'] = contratos_alerta
        return context

class ContratoCreateView(LoginRequiredMixin, MessageMixin, CreateView):
    model = Contrato
    template_name = 'form_contrato.html'
    form_class = ContratoForm
    success_url = reverse_lazy('core:lista_contratos')
    success_message = 'Contrato criado com sucesso!'

    def form_valid(self, form):
        contrato_service = ContratoService(self.request.user)
        
        try:
            # Preparar dados do formulário
            dados = form.cleaned_data
            
            # Criar contrato usando service
            contrato = contrato_service.criar_contrato(dados)
            
            # Definir o objeto criado para redirecionamento
            self.object = contrato
            
            messages.success(self.request, self.success_message)
            return redirect(self.success_url)
            
        except ValidationError as e:
            form.add_error(None, str(e))
            return self.form_invalid(form)
        except Exception as e:
            messages.error(self.request, f'Erro ao criar contrato: {str(e)}')
            return self.form_invalid(form)

class ContratoUpdateView(LoginRequiredMixin, MessageMixin, UpdateView):
    model = Contrato
    template_name = 'form_contrato.html'
    form_class = ContratoForm
    success_url = reverse_lazy('core:lista_contratos')
    success_message = 'Contrato atualizado com sucesso!'

    def form_valid(self, form):
        contrato_service = ContratoService(self.request.user)
        
        try:
            # Preparar dados do formulário
            dados = form.cleaned_data
            
            # Atualizar contrato usando service
            contrato = contrato_service.atualizar_contrato(self.object.id, dados)
            
            messages.success(self.request, self.success_message)
            return redirect(self.success_url)
            
        except ValidationError as e:
            form.add_error(None, str(e))
            return self.form_invalid(form)
        except Exception as e:
            messages.error(self.request, f'Erro ao atualizar contrato: {str(e)}')
            return self.form_invalid(form)

class ContratoDeleteView(LoginRequiredMixin, OwnerRequiredMixin, DeleteView):
    model = Contrato
    template_name = 'confirm_delete.html'
    success_url = reverse_lazy('core:lista_contratos')

    def delete(self, request, *args, **kwargs):
        try:
            response = super().delete(request, *args, **kwargs)
            messages.success(self.request, 'Contrato excluído com sucesso!')
            return response
        except Exception as e:
            messages.error(self.request, f'Erro ao excluir contrato: {str(e)}')
            return redirect(self.success_url)

# Views de Notas
class NotaListView(LoginRequiredMixin, FilterMixin, SearchMixin, PaginationMixin, ListView):
    model = Nota
    template_name = 'lista_notas.html'
    context_object_name = 'notas'
    paginate_by = 15
    paginator_class = OptimizedPaginator

    def get_queryset(self):
        nota_service = NotaService(self.request.user)
        
        # Obter filtros da requisição
        filtros = {
            'empresa': self.request.GET.get('empresa'),
            'setor': self.request.GET.get('setor'),
            'status': self.request.GET.get('status'),
            'data_inicio': self.request.GET.get('data_inicio'),
            'data_fim': self.request.GET.get('data_fim'),
            'contrato_id': self.request.GET.get('contrato_id')
        }
        
        # Remover filtros vazios
        filtros = {k: v for k, v in filtros.items() if v}
        
        return nota_service.listar_notas(filtros)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Adicionar empresas do cache para o filtro
        context['empresas'] = CacheManager.get_empresas_list()
        # Adicionar contratos para filtro
        contrato_service = ContratoService(self.request.user)
        context['contratos'] = contrato_service.listar_contratos()
        return context

class NotaCreateView(LoginRequiredMixin, MessageMixin, CreateView):
    model = Nota
    template_name = 'form_nota.html'
    form_class = NotaForm
    success_url = reverse_lazy('core:lista_notas')
    success_message = 'Nota criada com sucesso!'

    def form_valid(self, form):
        nota_service = NotaService(self.request.user)
        
        try:
            # Preparar dados do formulário
            dados = form.cleaned_data
            
            # Criar nota usando service
            nota = nota_service.criar_nota(dados)
            
            # Definir o objeto criado para redirecionamento
            self.object = nota
            
            messages.success(self.request, self.success_message)
            return redirect(self.success_url)
            
        except ValidationError as e:
            form.add_error(None, str(e))
            return self.form_invalid(form)
        except Exception as e:
            messages.error(self.request, f'Erro ao criar nota: {str(e)}')
            return self.form_invalid(form)

class NotaUpdateView(LoginRequiredMixin, MessageMixin, UpdateView):
    model = Nota
    template_name = 'form_nota.html'
    form_class = NotaForm
    success_url = reverse_lazy('core:lista_notas')
    success_message = 'Nota atualizada com sucesso!'

    def form_valid(self, form):
        try:
            response = super().form_valid(form)
            messages.success(self.request, self.success_message)
            return response
        except Exception as e:
            messages.error(self.request, f'Erro ao atualizar nota: {str(e)}')
            return self.form_invalid(form)

class NotaDeleteView(LoginRequiredMixin, OwnerRequiredMixin, DeleteView):
    model = Nota
    template_name = 'confirm_delete.html'
    success_url = reverse_lazy('core:lista_notas')

    def post(self, request, *args, **kwargs):
        if "cancel" in request.POST:
            return redirect('core:lista_notas')
        
        try:
            response = super().post(request, *args, **kwargs)
            messages.success(self.request, 'Nota excluída com sucesso!')
            return response
        except Exception as e:
            messages.error(self.request, f'Erro ao excluir nota: {str(e)}')
            return redirect(self.success_url)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['tipo_objeto'] = 'nota'
        return context

class ProcessarNotaView(LoginRequiredMixin, View):
    """View para processar (marcar como saída) uma nota"""
    
    def post(self, request, pk):
        nota_service = NotaService(request.user)
        
        try:
            data_saida = request.POST.get('data_saida')
            nota = nota_service.processar_nota(pk, data_saida)
            
            messages.success(request, f'Nota {nota.numero} processada com sucesso!')
            
        except ValidationError as e:
            messages.error(request, str(e))
        except Exception as e:
            messages.error(request, f'Erro ao processar nota: {str(e)}')
        
        return redirect('core:lista_notas')

# Views de Relatórios
class RelatoriosView(LoginRequiredMixin, TemplateView):
    template_name = 'relatorios.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        
        # Obter services
        relatorio_service = RelatorioService(self.request.user)
        contrato_service = ContratoService(self.request.user)
        
        # Obter filtros
        filtros = {
            'contrato_id': self.request.GET.get('contrato'),
            'data_inicio': self.request.GET.get('data_inicio'),
            'data_fim': self.request.GET.get('data_fim'),
            'empresa': self.request.GET.get('empresa'),
            'setor': self.request.GET.get('setor')
        }
        
        # Remover filtros vazios
        filtros = {k: v for k, v in filtros.items() if v}
        
        try:
            # Gerar relatório de notas
            relatorio_notas = relatorio_service.gerar_relatorio_notas(filtros)
            context['notas'] = relatorio_notas['notas'][:10]  # Limitar para exibição
            context['estatisticas_notas'] = relatorio_notas['estatisticas']
            
            # Gerar relatório de contratos
            relatorio_contratos = relatorio_service.gerar_relatorio_contratos(filtros)
            context['estatisticas_contratos'] = relatorio_contratos['estatisticas']
            
            # Adicionar contratos para filtro
            context['contratos'] = contrato_service.listar_contratos()
            
        except Exception as e:
            messages.error(self.request, f'Erro ao gerar relatório: {str(e)}')
            context['notas'] = []
            context['estatisticas_notas'] = {}
            context['estatisticas_contratos'] = {}
            context['contratos'] = []
        
        return context

    def post(self, request, *args, **kwargs):
        export_type = request.POST.get('export_type')
        relatorio_service = RelatorioService(request.user)
        
        # Obter filtros
        filtros = {
            'contrato_id': request.POST.get('contrato'),
            'data_inicio': request.POST.get('data_inicio'),
            'data_fim': request.POST.get('data_fim'),
            'empresa': request.POST.get('empresa'),
            'setor': request.POST.get('setor')
        }
        
        # Remover filtros vazios
        filtros = {k: v for k, v in filtros.items() if v}
        
        try:
            if export_type == 'pdf':
                # TODO: Implementar exportação para PDF
                messages.info(request, 'Exportação para PDF será implementada em breve')
                return redirect('core:relatorios')
            
            elif export_type == 'excel':
                # TODO: Implementar exportação para Excel
                messages.info(request, 'Exportação para Excel será implementada em breve')
                return redirect('core:relatorios')
            
            elif export_type == 'word':
                # TODO: Implementar exportação para Word
                messages.info(request, 'Exportação para Word será implementada em breve')
                return redirect('core:relatorios')
                
        except Exception as e:
            messages.error(request, f'Erro na exportação: {str(e)}')
        
        return redirect('core:relatorios')

class GerarProtocoloView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        # Falha controlada caso python-docx não esteja disponível (como em Vercel)
        if not DOCX_AVAILABLE:
            messages.error(request, 'Geração de protocolo indisponível: dependência python-docx ausente no ambiente.')
            return redirect('core:relatorios')
        notas_ids = request.POST.getlist('notas_selecionadas')
        notas = Nota.objects.filter(id__in=notas_ids).order_by('data_entrada')
        despacho_numero = request.POST.get('despacho_numero', '').strip()
        secretaria_nome = request.POST.get('secretaria', '').strip()
        
        if not notas:
            messages.error(request, 'Nenhuma nota selecionada.')
            return redirect('core:home')

        # Criar documento Word
        doc = Document()
        
        # Configurar margens
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2)
            section.right_margin = Cm(2)

        # Padronização de fontes e espaçamentos (Times 12, espaçamento simples)
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = 'Times New Roman'
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            normal_style.font.size = Pt(12)
        except Exception:
            pass

        # Funções auxiliares
        def add_paragraph(text, bold=False, align=None, size=12, color=None, space_after=Pt(6)):
            p = doc.add_paragraph(text)
            run = p.runs[0]
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color
            if align is not None:
                p.alignment = align
            p.paragraph_format.space_after = space_after
            return p

        def nome_mes_portugues(m):
            meses = [
                'janeiro','fevereiro','março','abril','maio','junho',
                'julho','agosto','setembro','outubro','novembro','dezembro'
            ]
            return meses[m-1]

        # Variáveis dinâmicas de cabeçalho/assinatura a partir do ambiente
        municipio = os.getenv('MUNICIPIO_NOME', 'Castanhal (PA)')
        coordenador_nome = os.getenv('COORDENADOR_NOME', 'Helton J. de S. Trajano da S. Teles')
        portaria_numero = os.getenv('PORTARIA_NUMERO', 'Portaria nº 279/2025')
        contador_nome = os.getenv('CONTADOR_NOME', 'Contador')
        contador_crc = os.getenv('CONTADOR_CRC', 'CRC/XXXXXXXXXXXX')

        # Cabeçalho institucional apenas com texto centralizado
        header_text = (
            'ESTADO DO PARÁ\nPREFEITURA MUNICIPAL DE CASTANHAL\nCOORDENADORIA DE CONTROLE INTERNO\n'
            'E-MAIL: coordenadoriacontrolcastanhal@gmail.com'
        )
        p_hdr = doc.add_paragraph('')
        r_hdr = p_hdr.add_run(header_text)
        r_hdr.font.size = Pt(10)
        r_hdr.font.bold = True
        p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('')  # espaçamento após cabeçalho

        # Título principal
        ano_atual = datetime.now().strftime('%Y')
        # Título com número do despacho informado
        numero_despacho_texto = despacho_numero if despacho_numero else f'XXX/{ano_atual}'
        titulo_texto = f'DESPACHO DO CONTROLE INTERNO Nº {numero_despacho_texto}'
        add_paragraph(titulo_texto, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, space_after=Pt(12))

        # Vocativo e Secretaria
        add_paragraph('Ilmo. Sr.º', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=Pt(2))
        secretaria_texto = secretaria_nome if secretaria_nome else 'Secretaria Municipal de Assistência Social,'
        p_sec = doc.add_paragraph('')
        r_sec = p_sec.add_run(secretaria_texto)
        r_sec.font.size = Pt(12)
        # Cor preta conforme solicitado
        r_sec.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        r_sec.underline = True
        p_sec.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_sec.paragraph_format.space_after = Pt(12)

        # Texto introdutório conforme modelo
        intro = (
            'Em atenção à solicitação de manifestação, desta Coordenadoria de Controle Interno do Município quanto à legalidade '
            'do pagamento da referida despesa, em conformidade com os critérios das Ordens de Fornecimentos e, fundamentando-se '
            'nos artigos 62 e seguintes da Lei nº 4.320/64 que discorre sobre pagamento de despesas, assim como a sua regular '
            'liquidação, manifesta-se que:'
        )
        add_paragraph(intro, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=Pt(12))

        # Seção 1 - Do relatório
        add_paragraph('1. DO RELATÓRIO', bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, size=12, space_after=Pt(6))
        corpo = (
            'Este Controle Interno promoveu a análise documental quanto ao prosseguimento da despesa abaixo descrita, considerando '
            'a documentação acostada, a verificação do direito adquirido pelo(s) prestador(es), tendo por base os títulos e documentos '
            'comprobatórios do respectivo crédito, apurando a origem, o objeto, a importância exata a ser paga, constatando o cumprimento '
            'do objeto através da juntada da Nota Fiscal (devidamente atestada pelos fiscais do contrato e Notas de Empenho), cabendo assim '
            'o prosseguimento do feito e efetivo pagamento visto o cumprimento dos ditames legais.'
        )
        add_paragraph(corpo, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=Pt(12))

        # Título da planilha
        add_paragraph('PLANILHA DE PAGAMENTO:', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12, space_after=Pt(6))

        # Criar tabela de acordo com novo modelo
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'

        header_cells = table.rows[0].cells
        headers = ['EMPENHO', 'EMPRESA', 'UNID. ORÇ', 'NF', 'DATA NF', 'VALOR']
        for i, text in enumerate(headers):
            p = header_cells[i].paragraphs[0]
            run = p.add_run(text)
            run.bold = True
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Adicionar dados dinâmicos
        for nota in notas:
            row_cells = table.add_row().cells
            row_cells[0].text = nota.empenho or '-'
            row_cells[1].text = nota.empresa
            row_cells[2].text = nota.setor or '-'
            row_cells[3].text = str(nota.numero)
            row_cells[4].text = nota.data_nota.strftime('%d/%m/%Y')
            row_cells[5].text = f'R$ {nota.valor:,.2f}'.replace(',', '.')

        # Espaço antes de data e assinaturas
        doc.add_paragraph('')

        # Data local por extenso
        hoje = datetime.now()
        data_extenso = f"{municipio}, {hoje.day:02d} de {nome_mes_portugues(hoje.month)} de {hoje.year}."
        add_paragraph(data_extenso, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT, size=12, space_after=Pt(12))

        # Assinaturas (duas colunas)
        assinatura_table = doc.add_table(rows=2, cols=2)
        assinatura_table.autofit = True
        # Linhas de assinatura
        for c in assinatura_table.rows[0].cells:
            p = c.paragraphs[0]
            r = p.add_run('_' * 40)
            r.font.size = Pt(12)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Nomes e cargos centralizados com fonte 10
        left_info = f"{coordenador_nome}\nCoordenador de Controle Interno do Município\n{portaria_numero}"
        right_info = f"{contador_nome}\nContador\n{contador_crc}"
        for idx, info in enumerate([left_info, right_info]):
            cell = assinatura_table.rows[1].cells[idx]
            # Limpa e adiciona parágrafo formatado
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(info)
            r.font.size = Pt(10)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Rodapé de recebimento
        doc.add_paragraph('')
        add_paragraph('RECEBIMENTO PELA SECRETARIA', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, space_after=Pt(0))

        # Salvar documento
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)

        # Gerar nome do arquivo: despacho + data e hora
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f'despacho_{timestamp}.docx'

        response = HttpResponse(
            f,
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename={filename}'
        
        return response

class ProtocoloPreviewView(LoginRequiredMixin, View):
    def post(self, request, *args, **kwargs):
        try:
            notas_ids = request.POST.getlist('notas_selecionadas')
            despacho_numero = request.POST.get('despacho_numero', '').strip()
            secretaria_nome = request.POST.get('secretaria', '').strip()

            if not notas_ids:
                return JsonResponse({'success': False, 'error': 'Nenhuma nota selecionada.'})

            notas = Nota.objects.filter(id__in=notas_ids).order_by('data_entrada')
            if not notas.exists():
                return JsonResponse({'success': False, 'error': 'Notas não encontradas.'})

            # Variáveis de ambiente usadas no documento
            municipio = os.getenv('MUNICIPIO_NOME', 'Castanhal (PA)')
            coordenador_nome = os.getenv('COORDENADOR_NOME', 'Helton J. de S. Trajano da S. Teles')
            portaria_numero = os.getenv('PORTARIA_NUMERO', 'Portaria nº275/2025')
            contador_nome = os.getenv('CONTADOR_NOME', 'Contador')
            contador_crc = os.getenv('CONTADOR_CRC', 'CRC/XXXXXXXXXXXX')

            # Data por extenso
            def nome_mes_portugues(m):
                meses = ['janeiro','fevereiro','março','abril','maio','junho','julho','agosto','setembro','outubro','novembro','dezembro']
                return meses[m-1]

            hoje = datetime.now()
            data_extenso = f"{municipio}, {hoje.day:02d} de {nome_mes_portugues(hoje.month)} de {hoje.year}."
            ano_atual = hoje.strftime('%Y')
            numero_despacho_texto = despacho_numero if despacho_numero else f'XXX/{ano_atual}'

            context = {
                'notas': notas,
                'despacho_numero': numero_despacho_texto,
                'secretaria_nome': secretaria_nome or 'Secretaria Municipal de Assistência Social,',
                'municipio': municipio,
                'data_extenso': data_extenso,
                'coordenador_nome': coordenador_nome,
                'portaria_numero': portaria_numero,
                'contador_nome': contador_nome,
                'contador_crc': contador_crc,
            }

            html = render_to_string('partials/protocolo_preview.html', context, request=request)
            return JsonResponse({'success': True, 'html': html})
        except Exception as e:
            return JsonResponse({'success': False, 'error': str(e)})

# Views de Usuário
class UsuarioListView(LoginRequiredMixin, AdminRequiredMixin, SearchMixin, PaginationMixin, ListView):
    model = Usuario
    template_name = 'lista_usuarios.html'
    context_object_name = 'usuarios'
    paginate_by = 10

    def get_queryset(self):
        usuario_service = UsuarioService(self.request.user)
        
        # Obter filtros da requisição
        filtros = {
            'ativo': self.request.GET.get('ativo'),
            'staff': self.request.GET.get('staff'),
            'busca': self.request.GET.get('busca')
        }
        
        # Converter strings para boolean quando necessário
        if filtros['ativo'] is not None:
            filtros['ativo'] = filtros['ativo'].lower() == 'true'
        if filtros['staff'] is not None:
            filtros['staff'] = filtros['staff'].lower() == 'true'
        
        # Remover filtros vazios
        filtros = {k: v for k, v in filtros.items() if v is not None and v != ''}
        
        return usuario_service.listar_usuarios(filtros)

class UsuarioCreateView(LoginRequiredMixin, AdminRequiredMixin, MessageMixin, CreateView):
    model = Usuario
    template_name = 'form_usuario.html'
    form_class = UsuarioForm
    success_url = reverse_lazy('core:lista_usuarios')
    success_message = 'Usuário criado com sucesso!'

    def form_valid(self, form):
        usuario_service = UsuarioService(self.request.user)
        
        try:
            # Preparar dados do formulário
            dados = form.cleaned_data
            
            # Criar usuário usando service
            usuario = usuario_service.criar_usuario(dados)
            
            # Definir o objeto criado para redirecionamento
            self.object = usuario
            
            messages.success(self.request, self.success_message)
            return redirect(self.success_url)
            
        except ValidationError as e:
            form.add_error(None, str(e))
            return self.form_invalid(form)
        except Exception as e:
            messages.error(self.request, f'Erro ao criar usuário: {str(e)}')
            return self.form_invalid(form)

class UsuarioUpdateView(LoginRequiredMixin, AdminRequiredMixin, MessageMixin, UpdateView):
    model = Usuario
    template_name = 'form_usuario.html'
    form_class = UsuarioUpdateForm
    success_url = reverse_lazy('core:lista_usuarios')
    success_message = 'Usuário atualizado com sucesso!'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['is_update'] = True
        return context

    def form_valid(self, form):
        usuario_service = UsuarioService(self.request.user)
        
        try:
            # Preparar dados do formulário
            dados = form.cleaned_data
            
            # Atualizar usuário usando service
            usuario = usuario_service.atualizar_usuario(self.object.id, dados)
            
            messages.success(self.request, self.success_message)
            return redirect(self.success_url)
            
        except ValidationError as e:
            form.add_error(None, str(e))
            return self.form_invalid(form)
        except Exception as e:
            messages.error(self.request, f'Erro ao atualizar usuário: {str(e)}')
            return self.form_invalid(form)
